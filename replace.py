import os
import sys
from glob import glob
from docx import Document

# usage:
# python test_tb.py C:\\Users\\童\\Desktop\\dat1 LEVEL2 LEVEL1

def th_table(folder_path, old_string, new_string):
    # find all doc files in the folder
    for filename in glob(os.path.join(folder_path, '*.docx')):
    # 打开文档
        document = Document(filename)
        # 遍历所有表格的单元格
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 如果只是为了内容，直接替换cell.text,但是为了保存原有格式，需要将每个单元格的文本当作一段看待，以此提取出run来不修改原格式
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_string in run.text:
                                    run.text = run.text.replace(old_string, new_string)
        # 保存文档
        document.save(filename)

def th_document(folder_path, old_string, new_string):
    # find all doc files in the folder
    for filename in glob(os.path.join(folder_path, '*.docx')):
        # open the doc file
        doc = Document(filename)
        for p in doc.paragraphs:
            if old_string in p.text:
                inline = p.runs
                # replace old_string with new_string in the runs
                for i in range(len(inline)):
                    if old_string in inline[i].text:
                        text = inline[i].text.replace(old_string, new_string)
                        inline[i].text = text
        doc.save(filename)

def main():
    folder_path = sys.argv[1]
    old_string = sys.argv[2]
    new_string = sys.argv[3]
    th_document(folder_path, old_string, new_string)
    th_table(folder_path, old_string, new_string)

if __name__ == "__main__":
    main()
